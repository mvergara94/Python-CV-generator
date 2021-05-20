
#pip3 install -r requirements.txt
#run on terminal
import docx
from docx import Document
from docx.shared import Inches
from data_validacao import DataBr

document = Document()

# profile picture
document.add_picture('photo.png', width=Inches(2.0))

# name, phone, and email details
name = input('What is your name ?')
phone_number = input('What is your phone number ?')
email = input('What is your e-mail ?')
marital_state = input('What is your marital state?')
adress = input('What is your adress ?')
document.add_paragraph(f'{name} | {phone_number} | {email}')
document.add_paragraph(f'{marital_state} | {adress}')

#about me
about_me_confirmation = input("Do you wanna add a 'About me' Section ?Yes or No")
if about_me_confirmation.lower() == 'yes':
    document.add_heading('About me')
    about_me = input('Tell about yourself:')
    document.add_paragraph(about_me)

#Formation
document.add_heading('Education and Certification')
while True:
    formation = input("Do you wish to add a formation or certification ? Yes or No")
    if formation.lower() == 'yes':
        e = document.add_paragraph()
        formacao = input("what is the formation or name of the course ?")
        instituicao = input("Name of the school or College ?")
        from_date = input('From Data')
        to_date = input('To Date')
        experience_details = input(f'Describe you experience at {instituicao}')
        e.add_run(f'{formacao}').bold = True
        e.add_run(f'- {instituicao}\n')
        e.add_run(f'{from_date} - {to_date} \n').italic = True
        e.add_run(experience_details)
    else:
        break


#Skills
document.add_heading('Skills')
while True:
    has_skill = input('Do you want to add a list of skills ? Yes or No ')
    if has_skill.lower() == 'yes':
        skill = input('Enter skill')
        p = document.add_paragraph(skill)
        p.style = 'List Bullet'
    else:
        break


#work experience
document.add_heading('Work Experience')
while True:
    has_experiences = input('Do you want to add a job experience ? Yes or No ')
    if has_experiences.lower() == 'yes':
        p = document.add_paragraph()
        company = input('Enter company')
        from_date = input('From Data')
        to_date = input('To Date')
        p.add_run(f'{company} ').bold = True
        p.add_run(f'{from_date} - {to_date} \n').italic = True
        experience_detaisl = input(f'Describe you experience at {company}')
        p.add_run(experience_detaisl)
    else:
        break

#footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
data = DataBr()
p.text = f'{data}'

document.save('cv.docx')




